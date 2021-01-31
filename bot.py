# Load libraries
# ... Utility libraries
from uuid import uuid1
from docx import Document
import os
import re
import random
# ... Telegram bot libraries 
from telegram import Update
from telegram.ext import Updater, CommandHandler, MessageHandler, Filters, CallbackContext
# ... model libraries 
from transformers import AutoModelWithLMHead, AutoTokenizer
import json
from flair.models import SequenceTagger
from flair.data import Sentence
import json
from random import seed, randint
from nltk import pos_tag
import nltk
nltk.download('averaged_perceptron_tagger')
nltk.download('stopwords')
nltk.download('universal_tagset')
import pke

# We use a pretrained question generation model (https://huggingface.co, "mrm8488/t5-base-finetuned-question-generation-ap") and a pretrained NER model (flair)
tokenizer = AutoTokenizer.from_pretrained("mrm8488/t5-base-finetuned-question-generation-ap")
CGmodel = AutoModelWithLMHead.from_pretrained("mrm8488/t5-base-finetuned-question-generation-ap")
NERmodel = SequenceTagger.load('ner-ontonotes-fast') #.load('ner')

#----------------------------------------------------------------------------------------
# Functions for model logic
#----------------------------------------------------------------------------------------
def get_question(answer, context, max_length=64):
  input_text = "answer: %s  context: %s </s>" % (answer, context)
  features = tokenizer([input_text], return_tensors='pt')

  output = CGmodel.generate(input_ids=features['input_ids'], 
               attention_mask=features['attention_mask'],
               max_length=max_length)
  #question = tokenizer.decode(output[0])[len('<pad> question: '):-1 * len('</s>')]question: 
  question = tokenizer.decode(output[0])[len('question: '):]
  return question


def do_NER(context):
  s = Sentence(context)
  NERmodel.predict(s)
  raw = s.to_dict(tag_type='ner')
  answers = []
  for item in raw['entities']:
      answers.append(item['text'])
  if not answers:
    answers = get_key_words(context)
  return list(map(lambda x: x.capitalize(), list(set(map(lambda x: x.lower(), answers)))))


def get_questions(answers, context):
  generated_questions = []
  for answer in answers:
    #print('context:', context, '\n', 'answer:', answer, '\n', 'generated question:', get_question(answer, context))
    generated_questions.append({'question':get_question(answer, context), 'answer':answer})
  return generated_questions

def generate_files_by_message(text):
  dict_of_questions = get_questions(do_NER(text), text)
  return save_questions_into_docx_files(text, dict_of_questions)

def get_key_words(text, num_nouns=3):
  pos = {'NOUN', 'PROPN', 'ADJ'}
  extractor = pke.unsupervised.SingleRank()
  extractor.load_document(input=text,
                          language='en',
                          normalization=None)
  extractor.candidate_selection(pos=pos)
  extractor.candidate_weighting(window=1,
                                pos=pos)
  keyphrases = extractor.get_n_best(n=10)
  return random.choices(list(map(lambda x: x[0], keyphrases)), k=num_nouns)  

def get_human_readable_file_names(text, file_extantion = '.docx', num_tokens_in_filename=5):
    tokens = [token for token in re.split('[^A-Za-z0-9]', text) if token]
    normalized_beginning = '_'.join(tokens[:num_tokens_in_filename])
    student_file_name = '_'.join([normalized_beginning, 'student']) + file_extantion
    teacher_file_name = '_'.join([normalized_beginning, 'teacher']) + file_extantion
    return student_file_name, teacher_file_name

def get_unique_file_names(text, file_extantion = '.docx'):
    file_id = uuid1().hex
    student_file_name = '_'.join([file_id, 'student']) + file_extantion
    teacher_file_name = '_'.join([file_id, 'teacher']) + file_extantion
    return student_file_name, teacher_file_name


def save_questions_into_docx_files(text:str, questions:list, upload_folder=''):
    '''
    Функция принимает текст и вопросы к нему, формирует и сохраняет два документа .docx (l)
    и возвращает пути к ним
    
    questions - данные в формате
    [
      { 
        'answer': 'the last few years',
        'question': 'When did Deep Learning pick up pace in academia and industry?'
      },
     {
       'answer': 'Machine Learning',
        'question': 'What type of practitioners have infiltrated the academia to its roots?'
     }
    ]
    '''
    student_docx_name, teacher_docx_name = get_unique_file_names(text)
    student_docx_path = os.path.join(upload_folder, student_docx_name)
    teacher_docx_path = os.path.join(upload_folder, teacher_docx_name)
    student_document = Document()
    teacher_document = Document()
    for document in [student_document, teacher_document]:
        document.add_heading('Text', 1)
        paragraphs = [paragraph for paragraph in text.split('\n') if paragraph]
        for paragraph in paragraphs:
            document.add_paragraph(paragraph)
    student_document.add_heading('Questions', 1)
    teacher_document.add_heading('Questions with answers', 1)
    for question_id, question_data in enumerate(questions):
        question = question_data['question']
        answer = question_data['answer']
        enumerated_question = str(question_id+1) + '. ' + question
        enumerated_question_with_answer = enumerated_question + '\nAnswer: ' + answer 
        student_document.add_paragraph(enumerated_question)
        teacher_document.add_paragraph(enumerated_question_with_answer)
    student_document.save(student_docx_path)
    teacher_document.save(teacher_docx_path)
    return student_docx_path, teacher_docx_path

#----------------------------------------------------------------------------------------
# Functions for bot 
#----------------------------------------------------------------------------------------
"""
Simple Bot to reply to Telegram messages.
First, a few handler functions are defined. Then, those functions are passed to
the Dispatcher and registered at their respective places.
Then, the bot is started and runs until we press Ctrl-C on the command line.
Usage:
Basic Echobot example, repeats messages.
Press Ctrl-C on the command line or send a signal to the process to stop the
bot.
"""
# Load bot token from file. Get your own from @botfather https://t.me/botfather
with open('token.txt', 'r') as txt:
  for line in txt:
    token = line

# Define a few command handlers. These usually take the two arguments update and
# context. Error handlers also receive the raised TelegramError object in error.
def start(update: Update, context: CallbackContext) -> None:
    """Send a message when the command /start is issued."""
    start_mes = '''Привет, я QuizGuy, чтобы составить тест в автоматическом режиме просто пришли мне текст (На английском!) в следующем сообщении!
    '''
    update.message.reply_text(start_mes)

def help(update: Update, context: CallbackContext) -> None:
    """Send a message when the command /start is issued."""
    help_mes = '''Чтобы сгенерировать тест в автоматическом режиме, просто пришли мне текст (на английском) в следующем сообщении.
    '''
    update.message.reply_text(help_mes)


def chat(update: Update, context: CallbackContext) -> None:
    """Chat with user """
    """
    # Logging user request for development purposes 
    print(f"Handling message  {update.message.chat_id}-{update.message.message_id} from {update.message.chat.username}")
    # Log messages to file
    with open(f"log/{update.message.chat_id}-{update.message.message_id}.log", "a", encoding="utf-8") as f:
        f.write(f",\n{update.message.to_json()}")
    """
    # Pass user message to bot_reply for processing 
    grateful_list = ['Спасибо, что воспользовались нашим ботом!', 'Будем рады составить для вас еще одно задание по тексту!', 'Надеемся, что наши вопросы вам понравятся.']
    user_mes = update.message.text
    context.bot.send_message(chat_id=update.effective_chat.id, text='Начинаю обработку нейросетями...')
    student_docx_path, teacher_docx_path = generate_files_by_message(user_mes)
    student_human_readable_filename, teacher_human_readable_filename = get_human_readable_file_names(user_mes)
    context.bot.send_message(chat_id=update.effective_chat.id, text='Подготавливаю файлы...')
    context.bot.send_document(chat_id=update.effective_chat.id, document=open(student_docx_path, 'rb'), filename=student_human_readable_filename)
    context.bot.send_document(chat_id=update.effective_chat.id, document=open(teacher_docx_path, 'rb'), filename=teacher_human_readable_filename)
    context.bot.send_message(chat_id=update.effective_chat.id, text=random.choice(grateful_list))

def main():
    """Start the bot."""
    # Create the Updater and pass it your bot's token.
    # Make sure to set use_context=True to use the new context based callbacks
    # Post version 12 this will no longer be necessary
    updater = Updater(token, use_context=True)

    # Get the dispatcher to register handlers
    dispatcher = updater.dispatcher

    # on different commands - answer in Telegram
    dispatcher.add_handler(CommandHandler("start", start))
    dispatcher.add_handler(CommandHandler("help", help))
    # on noncommand i.e message - echo the message on Telegram
    dispatcher.add_handler(MessageHandler(Filters.text & ~Filters.command, chat))

    # Start the Bot
    updater.start_polling()

    # Run the bot until you press Ctrl-C or the process receives SIGINT,
    # SIGTERM or SIGABRT. This should be used most of the time, since
    # start_polling() is non-blocking and will stop the bot gracefully.
    updater.idle()

if __name__ == '__main__':
    main()    